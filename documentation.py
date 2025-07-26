import yfinance as yf
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from datetime import datetime
import streamlit as st
from io import BytesIO

# --- 1. User Inputs ---
TICKER = "NEE"
ANALYSIS_DATE = "2024-12-31"
REPORT_TEMPLATE = "your_template.docx"

st.title("Stock Analysis Report Generator")

# --- 2. Fetch Data ---
stock = yf.Ticker(TICKER)
hist = stock.history(period="1y")
info = stock.info

# --- 3. Generate Price Chart ---
plt.figure(figsize=(6,3))
plt.plot(hist.index, hist['Close'])
plt.title(f"{TICKER} Price (1 Year)")
plt.xlabel("Date")
plt.ylabel("Close Price")
plt.tight_layout()
plt.savefig("price_chart.png")
plt.close()

# --- 4. Start Report from Template ---
doc = Document(REPORT_TEMPLATE)

# --- 5. Fill Executive Summary (Page 1) ---
doc.paragraphs[0].text = f"Stock Analysis Report: {info['shortName']} ({TICKER})"
doc.add_paragraph(f"Exchange: {info.get('exchange', 'N/A')}")
doc.add_paragraph(f"Analysis Date: {ANALYSIS_DATE}")
doc.add_paragraph(f"Current Price: ${hist['Close'][-1]:.2f}")

# Example: Insert chart
doc.add_picture("price_chart.png", width=Inches(4.5))

# --- 6. Add Business Description ---
doc.add_heading("Business Description", level=1)
doc.add_paragraph(info.get('longBusinessSummary', 'No summary available.'))

# --- 7. Add Financials Table Example ---
doc.add_heading("Key Financials", level=1)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'
hdr_cells[2].text = 'Year'
table.add_row().cells[0].text = 'Revenue'
table.rows[1].cells[1].text = str(info.get('totalRevenue', 'N/A'))
table.rows[1].cells[2].text = str(datetime.now().year)

# --- 7b. Add Valuation Section ---
doc.add_heading("Valuation", level=1)

# --- DDM Calculation (Dividend Discount Model) ---
dividend = info.get('dividendRate', 0)  # Annual dividend per share
div_growth = 0.05  # 5% growth rate (example)
cost_of_equity = 0.08  # 8% required return (example)

if dividend > 0 and cost_of_equity > div_growth:
    ddm_value = dividend * (1 + div_growth) / (cost_of_equity - div_growth)
    doc.add_paragraph(f"DDM Valuation: ${ddm_value:.2f} per share (Dividend: ${dividend:.2f}, Growth: {div_growth*100:.1f}%, Cost of Equity: {cost_of_equity*100:.1f}%)")
else:
    doc.add_paragraph("DDM Valuation: Not applicable (missing dividend or invalid growth/discount rates)")

# --- DCF Calculation (Discounted Cash Flow, simple version) ---
fcf = info.get('freeCashflow', 0)  # Most recent free cash flow
fcf_growth = 0.04  # 4% growth rate (example)
discount_rate = 0.08  # 8% discount rate (example)
terminal_growth = 0.03  # 3% terminal growth (example)
years = 5

if fcf and fcf > 0:
    dcf_sum = 0
    for t in range(1, years+1):
        dcf_sum += (fcf * (1 + fcf_growth) ** t) / ((1 + discount_rate) ** t)
    terminal_value = (fcf * (1 + fcf_growth) ** years) * (1 + terminal_growth) / ((discount_rate - terminal_growth) * (1 + discount_rate) ** years)
    dcf_value = dcf_sum + terminal_value
    shares_outstanding = info.get('sharesOutstanding', 1)
    dcf_per_share = dcf_value / shares_outstanding
    doc.add_paragraph(f"DCF Valuation: ${dcf_per_share:.2f} per share (FCF: ${fcf:.2f}, Growth: {fcf_growth*100:.1f}%, Discount Rate: {discount_rate*100:.1f}%)")
else:
    doc.add_paragraph("DCF Valuation: Not applicable (missing or invalid free cash flow data)")

# --- 8. Download Report ---
buffer = BytesIO()
doc.save(buffer)
buffer.seek(0)

st.success("Report generated!")
st.download_button(
    label="Download Word Report",
    data=buffer,
    file_name=f"Stock_Report_{TICKER}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


